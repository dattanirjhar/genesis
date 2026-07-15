"""
Reporter — the Markdown knowledge base into a DOCX assessment report.

Pipeline position (the bottom of the stack):

    knowledge/findings/*.md  ->  Reporter  ->  knowledge/reports/genesis_report.docx

It consumes the human-readable knowledge layer (source of truth for findings),
groups findings by severity, and writes a structured Word document. An optional
executive summary is drafted by the reasoner, so the report is grounded in the
same RAG pipeline as the analyst Q&A.

    python -m report.report          # deterministic report
    python -m report.report --llm    # + LLM-drafted executive summary
"""

from __future__ import annotations

import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from config import KNOWLEDGE_DIR, REPORTS_DIR  # noqa: E402

# Highest to lowest, for ordering findings and the summary table.
SEVERITY_ORDER = ["critical", "high", "medium", "low", "informational", "info", "unknown"]


def _split_frontmatter(markdown: str) -> tuple[dict, str]:
    """Split a finding document into its front-matter dict and body."""
    if not markdown.startswith("---"):
        return {}, markdown
    lines = markdown.splitlines()
    try:
        end = next(i for i in range(1, len(lines)) if lines[i].strip() == "---")
    except StopIteration:
        return {}, markdown
    meta = {}
    for line in lines[1:end]:
        if ":" in line:
            key, _, value = line.partition(":")
            meta[key.strip()] = value.strip()
    return meta, "\n".join(lines[end + 1:]).strip()


def load_findings() -> list[dict]:
    """Load every knowledge finding as {meta, body, file}."""
    items = []
    for path in sorted(KNOWLEDGE_DIR.glob("*.md")):
        meta, body = _split_frontmatter(path.read_text(encoding="utf-8"))
        items.append({"meta": meta, "body": body, "file": path.name})
    return items


def _severity_rank(severity: str | None) -> int:
    sev = (severity or "unknown").lower()
    return SEVERITY_ORDER.index(sev) if sev in SEVERITY_ORDER else len(SEVERITY_ORDER)


def _default_summary(findings: list[dict], counts: Counter) -> str:
    hosts = {f["meta"].get("host", "unknown") for f in findings}
    parts = [f"{counts[s]} {s}" for s in SEVERITY_ORDER if counts.get(s)]
    return (
        f"This assessment identified {len(findings)} findings across "
        f"{len(hosts)} hosts ({', '.join(parts)}). Findings are detailed below, "
        f"ordered by severity. Review the critical and high-severity items first."
    )


def _add_body(doc: Document, body: str) -> None:
    """Render a finding's Markdown body into the document (light conversion)."""
    in_fence = False
    for line in body.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            continue
        if not stripped:
            continue
        if stripped.startswith("## "):
            run = doc.add_paragraph().add_run(stripped[3:].strip())
            run.bold = True
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif in_fence:
            doc.add_paragraph(line, style="Intense Quote")
        else:
            doc.add_paragraph(stripped)


def build_report(exec_summary: str | None = None, out_path: Path | None = None) -> Path:
    """Build the DOCX report and return its path."""
    findings = load_findings()
    findings.sort(key=lambda f: _severity_rank(f["meta"].get("severity")))
    counts = Counter((f["meta"].get("severity") or "unknown").lower() for f in findings)

    doc = Document()
    doc.add_heading("Genesis Security Assessment Report", level=0)
    doc.add_paragraph(f"Generated {datetime.now():%Y-%m-%d %H:%M} by Genesis")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(exec_summary or _default_summary(findings, counts))

    doc.add_heading("Findings Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text, header[1].text = "Severity", "Count"
    for sev in SEVERITY_ORDER:
        if counts.get(sev):
            row = table.add_row().cells
            row[0].text, row[1].text = sev.title(), str(counts[sev])

    doc.add_heading("Detailed Findings", level=1)
    for f in findings:
        m = f["meta"]
        severity = (m.get("severity") or "unknown").upper()
        title = m.get("name") or m.get("finding_id", "Finding")
        doc.add_heading(f"[{severity}] {title}", level=2)
        meta_line = doc.add_paragraph()
        meta_line.add_run(
            f"Host: {m.get('host', '?')}  |  Service: {m.get('service', '?')}  |  "
            f"Scanner: {m.get('scanner', '?')}  |  CVE: {m.get('cve', 'none')}  |  "
            f"Validated: {m.get('validated', 'false')}"
        ).italic = True
        _add_body(doc, f["body"])

    out_path = out_path or (REPORTS_DIR / "genesis_report.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_markdown_report(exec_summary: str | None = None,
                          out_path: Path | None = None) -> Path:
    """Build a Markdown report (report.md) and return its path.

    Markdown is the primary format — human-readable now, convertible to DOCX
    later. The finding bodies are already Markdown, so their ## subsections are
    demoted to #### to sit correctly under each ### finding heading.
    """
    findings = load_findings()
    findings.sort(key=lambda f: _severity_rank(f["meta"].get("severity")))
    counts = Counter((f["meta"].get("severity") or "unknown").lower() for f in findings)

    lines = [
        "# Genesis Security Assessment Report",
        "",
        f"_Generated {datetime.now():%Y-%m-%d %H:%M} by Genesis_",
        "",
        "## Executive Summary",
        "",
        exec_summary or _default_summary(findings, counts),
        "",
        "## Findings Summary",
        "",
        "| Severity | Count |",
        "| --- | --- |",
    ]
    for sev in SEVERITY_ORDER:
        if counts.get(sev):
            lines.append(f"| {sev.title()} | {counts[sev]} |")
    lines += ["", "## Detailed Findings", ""]

    for f in findings:
        m = f["meta"]
        severity = (m.get("severity") or "unknown").upper()
        title = m.get("name") or m.get("finding_id", "Finding")
        lines.append(f"### [{severity}] {title}")
        lines.append("")
        lines.append(
            f"**Host:** {m.get('host', '?')}  |  **Service:** {m.get('service', '?')}"
            f"  |  **Scanner:** {m.get('scanner', '?')}  |  **CVE:** {m.get('cve', 'none')}"
            f"  |  **Validated:** {m.get('validated', 'false')}"
        )
        lines.append("")
        body = f["body"].replace("\n## ", "\n#### ")
        if body.startswith("## "):
            body = "#### " + body[3:]
        lines.append(body)
        lines.append("")

    out_path = out_path or (REPORTS_DIR / "genesis_report.md")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


def _llm_executive_summary() -> str | None:
    """Draft an executive summary via the reasoner (grounded in the RAG index)."""
    try:
        from llm.reasoner import answer
        return answer(
            "Write a concise executive summary of this penetration test: the "
            "overall risk posture, the most critical findings, and the top "
            "remediation priorities."
        )
    except Exception as exc:  # noqa: BLE001
        print(f"LLM executive summary skipped: {exc}")
        return None


if __name__ == "__main__":
    if not list(KNOWLEDGE_DIR.glob("*.md")):
        raise SystemExit(f"No knowledge to report on in {KNOWLEDGE_DIR}. "
                         "Build it first: python -m tools.genesis rebuild")

    summary = _llm_executive_summary() if "--llm" in sys.argv else None
    out = build_report(exec_summary=summary)
    print(f"Report written: {out}")
